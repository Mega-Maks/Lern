import docx

doc = docx.Document('C:\\Users\\Ryzhk\\PycharmProjects\\Учёба\\Нипигормаш_2\\Шаблон_dox.docx')
all_paras = doc.paragraphs
for para in all_paras:
    print(para.text)
    print("-------")


for table in doc.tables:
    print('Новая таблица---------------------------------------------------------------------------------------')
    for row in table.rows:
        string = ''
        print('Новый столбец------------------------------------------------------------------------------------------')
        for cell in row.cells:
            if cell.text == '((street_2))':
                cell.text = '           иСУЗДАЛЬь'

            print('Новая клетка:', cell.text)
doc.save('C:\\Users\\Ryzhk\\Desktop\\some_file9.docx')