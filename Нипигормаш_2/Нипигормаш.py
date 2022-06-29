import openpyxl
import os
import docx

wb = openpyxl.load_workbook('Шаблон_exel.xlsx')
ws = wb.active
document_number = int(input('введите номер документа '))
act_val = int(input('введите номер акта '))
doc = docx.Document('C:\\Users\\Ryzhk\\PycharmProjects\\Учёба\\Нипигормаш_2\\Шаблон_dox.docx')


def clear():
    os.chdir('C:\\Users\\Ryzhk\\Desktop\\Новая папка (2)')
    try:
        for i in os.listdir():
            os.remove(i)
    except PermissionError:
        print(f"Файл {i} открыт, закрой его")


class Docx:
    def __init__(self, way_list: list, date: str, doc_num: int, doc_tables):
        self.doc_tables = doc_tables
        self.doc_num = doc_num
        self.date = date
        self.street_1 = way_list[1]
        self.street_2 = way_list[2]
        self.money = way_list[3]
        self.tons = way_list[4]
        self.contain = way_list[5]
        self.tons_dict = {
            1: '1000 кг (Одна)',
            2: '2000 кг (Две)',
            3: '3000 кг (Три)',
            4: '4000 кг (Четыре)',
            5: '5000 кг (Пять)',
            6: '6000 кг (Шесть)',
            7: '7000 кг (Семь)',
            8: '8000 кг (Восемь)',
            9: '9000 кг (Девять)',
            10: '10000 кг (Десять)'
        }
        self.money_dict = {
            1500: '1500,0 (Одна тысяча пятьсот) руб 00коп',
            3000: '3000,0 (Три тысячи ) руб 00коп',
            4500: '4500,0 (Четыре тысячи пятьсот) руб 00коп',
            6000: '6000,0 (Шесть тысяч) руб 00коп',
            7500: '7500,0 (Семь тысяч пятьсот) руб 00коп',
            9000: '9000,0 (Девять тысяч ) руб 00коп',
            10500: '10500,0 (Девять тысяч пятьсот) руб 00коп',
            12000: '12000,0 (Десять тысяч ) руб 00коп',
            13500: '13500,0 (Тринадцать тысяч пятьсот) руб 00коп',
            15000: '15000,0 (Пятнадцать тысяч) руб 00коп',
            16500: '16500,0 (Шестнадцать тысяч пятьсот) руб 00коп',
            18000: '18000,0 (Восемнадцать тысяч ) руб 00коп',
            19500: '19500,0 (Девятнадцать тысяч пятьсот) руб 00коп',
            21000: '21000,0 (Двадцать одна тысяча) руб 00коп',
            22500: '22500,0 (Двадцать две тысячи пятьсот) руб 00коп'
        }

    def docx_writer(self):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text == '01.01.2020':
                        cell.text = self.date
                        cell.alignment = 2
                    elif cell.text == 'Ящики':
                        cell.text = self.contain
                    elif cell.text == '1000 кг (Одна)':
                        cell.text = self.tons_dict[int(self.tons)]
                        cell.alignment = 2
                    elif cell.text == '1000 кг':
                        cell.text = '             ' + self.tons_dict[int(self.tons)][0:8]
                        cell.alignment = 2
                    elif cell.text == '((тонны))':
                        cell.text = '                ' + self.tons
                    elif cell.text == '((street_1))':
                        cell.text = '                                ' + self.street_1
                    elif cell.text == '((street_2))':
                        cell.text = '                                ' + self.street_2
                    elif cell.text == '((Деньги))':
                        cell.text = self.money_dict[int(self.money)]
                    elif cell.text == '((дата))':
                        cell.text = '           ' + self.date
        doc.save('C:\\Users\\Ryzhk\\Desktop\\Новая папка (2)\\' + str(self.doc_num) + '      ' + self.date +
                 '         ' + self.street_1 + ' - ' + self.street_2 + ' ' + '.docx')


class Exel:

    def __init__(self, list_with_str: list, act_val: int, document_number: int):
        self.date = list_with_str[0][0]
        self.list_with_str = list_with_str
        self.act_val = act_val
        self.document_number = document_number
        self.month_dict = {
            1: 'Января 2022 г.',
            2: 'Февраля 2022 г.',
            3: 'Марта 2022 г.',
            4: 'Апреля 2022 г.',
            5: 'Мая 2022 г.',
            6: 'Июня 2022 г.',
            7: 'Июля 2022 г.',
            8: 'Августа 2022 г.',
            9: 'Сентября 2022 г.',
            10: 'Октября 2022 г.',
            11: 'Ноября 2022 г.',
            12: 'Декабря 2022 г.',
        }
        self.B_18_dict = {
            1500: 'Одна тысяча пятьсот рублей 00коп',
            3000: 'Три тысячи рублей 00коп',
            4500: 'Четыре тысячи пятьсот рублей 00коп',
            6000: 'Шесть тысяч рублей 00коп',
            7500: 'Семь тысяч пятьсот рублей 00коп',
            9000: 'Девять тысяч рублей 00коп',
            10500: 'Девять тысяч пятьсот рублей 00коп',
            12000: 'Десять тысяч рублей 00коп',
            13500: 'Тринадцать тысяч пятьсот рублей 00коп',
            15000: 'Пятнадцать тысяч рублей 00коп',
            16500: 'Шестнадцать тысяч пятьсот рублей 00коп',
            18000: 'Восемнадцать тысяч рублей 00коп',
            19500: 'Девятнадцать тысяч пятьсот рублей 00коп',
            21000: 'Двадцать одна тысяча рублей 00коп',
            22500: 'Двадцать две тысячи пятьсот рублей 00коп'
        }

    def exel_writer(self):
        list_with_streets = []
        street_string = ''
        money = 0
        for day_list in self.list_with_str:
            list_with_streets.append(day_list[1] + ' - ' + day_list[2] + ';\n')
            money += int(day_list[3])
        for street in list_with_streets:
            street_string += street
        ws['B3'] = "Акт № " + str(self.act_val) + " от " + self.date[3:5] + " " + \
                   self.month_dict[int(self.date[0:2])]
        ws['D12'] = self.date + 'г. Услуги манипулятора (МАРКА А/М MAN, Х 012 ХВ,' +\
                                'Шестернин Александр Олегович) по маршруту: \n' + street_string
        ws['U12'] = str(int(money / 1500))
        ws['Z12'] = str(1500)
        ws['AD12'] = str(money) + ',00'
        ws['B17'] = f"Всего оказано услуг 1 на сумму {money},00"
        ws['B18'] = self.B_18_dict[money]
        os.chdir('C:\\Users\\Ryzhk\\Desktop\\Новая папка (2)')
        wb.save(str(self.document_number) + ' ' + self.date + ' Акт №' + str(self.act_val) + '.xlsx')


if __name__ == "__main__":
    clear()
    os.chdir('C:\\Users\\Ryzhk\\PycharmProjects\\Учёба\\Нипигормаш_2')
    with open('input_file.txt') as input_file:
        input_file = input_file.read()

    input_file = [[i.split('\t') for i in day.split('\n')] for day in input_file.split('\n\n')]
    for i in input_file:
        print(i)
    for day_list in input_file:
        doc_count = len(day_list) + 1
        ex_val = Exel(day_list, act_val, document_number)
        ex_val.exel_writer()
        date = ex_val.date
        act_val += 1
        document_number += 1
        for way_list in day_list:
            doc = docx.Document('C:\\Users\\Ryzhk\\PycharmProjects\\Учёба\\Нипигормаш_2\\Шаблон_dox.docx')
            docx_file = Docx(way_list, date, document_number, doc.tables)
            docx_file.docx_writer()
            document_number += 1
